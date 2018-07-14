#!/usr/bin/env python
# -*- coding: utf-8 -*-
#author:zhl


##方法1 ，有报错，待排查
# import win32com
# import win32com.client
#
# def readWordFile(path):
#     ##调用系统word功能，可以处理doc和docx文件
#     mw=win32com.client.Dispatch("Word.Application")
#     #打开文件：
#     doc=mw.Documents.Open(path)
#
#     for paragraph in doc.Paragraphs:
#         line=paragraph.Range.Text
#         print(line)
#
#     ##关闭文件：
#     doc.Close()
#     ##退出word
#     mw.Quit()
#
#
# path=r'C:\Users\lin\PycharmProjects\python_study_1s\python_study\git-zhl\python-study\python 语法基础\d15_自动化办公与鼠标键盘模拟\4.word自动化读写\aa.doc'
#
# readWordFile(path)

#方法2；
# https://blog.csdn.net/woshisangsang/article/details/75221723

# Python可以利用python-docx模块处理word文档，处理方式是面向对象的。也就是说python-docx模块会把word文档，
# 文档中的段落、文本、字体等都看做对象，对对象进行处理就是对word文档的内容处理。

# 如果需要读取word文档中的文字（一般来说，程序也只需要认识word文档中的文字信息），需要先了解python-docx模块的几个概念。
#
# 1，Document对象，表示一个word文档。
# 2，Paragraph对象，表示word文档中的一个段落
# 3，Paragraph对象的text属性，表示段落中的文本内容。

#pip3 install python-docx

import docx

def readWord(path):
    #获取文档对象
    file=docx.Document(path)
    print("段落数："+str(len(file.paragraphs)))

    ##输出每一段的内容
    for para in file.paragraphs:
        print(para.text)

    ##输出段落编号鸡段落内容
    for i in range(len(file.paragraphs)):
        print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)



path=r'C:\Users\lin\PycharmProjects\python_study_1s\python_study\git-zhl\python-study\python 语法基础\d15_自动化办公与鼠标键盘模拟\4.word自动化读写\bb.docx'
readWord(path)

##备注说明，读取doc文件会报错，详细请参考 ：https://www.cnblogs.com/beikew/p/8001164.html
